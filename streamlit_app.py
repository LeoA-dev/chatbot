import streamlit as st
from openai import OpenAI
import os
import io
import base64
import textwrap

# Optional libraries for document/image parsing
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None
    canvas = None

# -----------------------------
# CONFIG
# -----------------------------
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
VECTOR_STORE_ID = st.secrets["OPENAI_VECTOR_STORE_ID"] # set this in your env variables

client = OpenAI(api_key=OPENAI_API_KEY)

MODEL_CHOICES = {
    "Standard - GPT-5.4": "gpt-5.4",
    "Mittel - GPT-5.5": "gpt-5.5",
    "Maximal - GPT-5.5 Pro": "gpt-5.5-pro",
}

EXPORT_CHOICES = {
    "Nicht speichern": None,
    "Als TXT speichern": "txt",
    "Als Word-Dokument speichern": "docx",
    "Als PDF speichern": "pdf",
}

WEB_SEARCH_CHOICES = {
    "Websuche aus": "off",
    "Websuche an": "on",
}

# -----------------------------
# STREAMLIT UI
# -----------------------------
st.set_page_config(page_title="RAG Chatbot", layout="centered")
st.title("CHECK-Ki Chatbot")

# File uploads: allow attaching photos and documents
uploaded_files = st.file_uploader("Fotos/Dokumente anhängen (optional)", type=None, accept_multiple_files=True)

def choose_model(model_label):
    return MODEL_CHOICES[model_label]


def get_value(item, key, default=None):
    if isinstance(item, dict):
        return item.get(key, default)

    return getattr(item, key, default)


def response_to_dict(response):
    if hasattr(response, "model_dump"):
        return response.model_dump()

    if hasattr(response, "to_dict"):
        return response.to_dict()

    return {}


def add_source(sources, seen_urls, url, title=None):
    if not url or url in seen_urls:
        return

    seen_urls.add(url)
    sources.append({
        "title": title or url,
        "url": url,
    })


def collect_annotation_sources(content_part, sources, seen_urls):
    for annotation in get_value(content_part, "annotations", []) or []:
        annotation_type = get_value(annotation, "type")
        if annotation_type == "url_citation":
            add_source(
                sources,
                seen_urls,
                get_value(annotation, "url"),
                get_value(annotation, "title"),
            )
            continue

        citation = get_value(annotation, "url_citation")
        if citation:
            add_source(
                sources,
                seen_urls,
                get_value(citation, "url"),
                get_value(citation, "title"),
            )


def collect_tool_sources(action, sources, seen_urls):
    for source in get_value(action, "sources", []) or []:
        add_source(
            sources,
            seen_urls,
            get_value(source, "url") or get_value(source, "source_url"),
            get_value(source, "title"),
        )


def collect_response_sources(response):
    sources = []
    seen_urls = set()

    for output_item in get_value(response, "output", []) or []:
        if get_value(output_item, "type") == "message":
            for content_part in get_value(output_item, "content", []) or []:
                collect_annotation_sources(content_part, sources, seen_urls)
        elif get_value(output_item, "type") == "web_search_call":
            collect_tool_sources(get_value(output_item, "action", {}), sources, seen_urls)

    response_data = response_to_dict(response)
    for output_item in response_data.get("output", []) or []:
        if output_item.get("type") == "message":
            for content_part in output_item.get("content", []) or []:
                collect_annotation_sources(content_part, sources, seen_urls)
        elif output_item.get("type") == "web_search_call":
            collect_tool_sources(output_item.get("action", {}), sources, seen_urls)

    return sources


def build_response_options(web_search_mode):
    if web_search_mode == "off":
        return {}

    return {
        "tools": [{"type": "web_search"}],
        "tool_choice": "auto",
        "include": ["web_search_call.action.sources"],
    }


def render_sources(sources):
    if not sources:
        return

    with st.expander("Quellen", expanded=True):
        for index, source in enumerate(sources, start=1):
            st.markdown(f"{index}. [{source['title']}]({source['url']})")


def build_conversation_text(messages):
    if not messages:
        return "Keine Nachrichten vorhanden."

    lines = []
    for message in messages:
        role = "User" if message["role"] == "user" else "Assistant"
        model = message.get("model")
        model_suffix = f" ({model})" if model else ""
        lines.append(f"{role}{model_suffix}:")
        lines.append(message["content"])
        if message.get("sources"):
            lines.append("Quellen:")
            for index, source in enumerate(message["sources"], start=1):
                lines.append(f"{index}. {source['title']} - {source['url']}")
        lines.append("")

    return "\n".join(lines).strip()


def build_docx_export(messages):
    if Document is None:
        return None

    doc = Document()
    doc.add_heading("CHECK-Ki Chatbot Conversation", level=1)
    for message in messages:
        role = "User" if message["role"] == "user" else "Assistant"
        model = message.get("model")
        heading = f"{role} ({model})" if model else role
        doc.add_heading(heading, level=2)
        doc.add_paragraph(message["content"])
        if message.get("sources"):
            doc.add_paragraph("Quellen:")
            for index, source in enumerate(message["sources"], start=1):
                doc.add_paragraph(f"{index}. {source['title']} - {source['url']}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def build_pdf_export(messages):
    if A4 is None or canvas is None:
        return None

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    left_margin = 50
    top_margin = height - 50
    line_height = 14
    y = top_margin

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(left_margin, y, "CHECK-Ki Chatbot Conversation")
    y -= line_height * 2

    for message in messages:
        role = "User" if message["role"] == "user" else "Assistant"
        model = message.get("model")
        heading = f"{role} ({model})" if model else role

        if y < 70:
            pdf.showPage()
            y = top_margin

        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawString(left_margin, y, heading)
        y -= line_height
        pdf.setFont("Helvetica", 10)

        for paragraph in message["content"].splitlines() or [""]:
            for line in textwrap.wrap(paragraph, width=95) or [""]:
                if y < 50:
                    pdf.showPage()
                    y = top_margin
                    pdf.setFont("Helvetica", 10)
                pdf.drawString(left_margin, y, line)
                y -= line_height
            y -= 4

        if message.get("sources"):
            pdf.setFont("Helvetica-Bold", 10)
            if y < 50:
                pdf.showPage()
                y = top_margin
            pdf.drawString(left_margin, y, "Quellen:")
            y -= line_height
            pdf.setFont("Helvetica", 9)
            for index, source in enumerate(message["sources"], start=1):
                source_text = f"{index}. {source['title']} - {source['url']}"
                for line in textwrap.wrap(source_text, width=105):
                    if y < 50:
                        pdf.showPage()
                        y = top_margin
                        pdf.setFont("Helvetica", 9)
                    pdf.drawString(left_margin, y, line)
                    y -= line_height

        y -= line_height

    pdf.save()
    output.seek(0)
    return output.getvalue()


def extract_text_from_file(uploaded_file):
    name = uploaded_file.name
    lower = name.lower()
    data = uploaded_file.getvalue()

    # Text-like files
    if lower.endswith(('.txt', '.md', '.csv')):
        try:
            return data.decode('utf-8')
        except Exception:
            return data.decode('latin-1', errors='ignore')

    # PDF
    if lower.endswith('.pdf') and PyPDF2 is not None:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            text = []
            for p in reader.pages:
                try:
                    text.append(p.extract_text() or "")
                except Exception:
                    continue
            return "\n\n".join(text)
        except Exception:
            return f"(Konnte PDF {name} nicht extrahieren.)"

    # Images: try OCR if available
    if lower.endswith(('.png', '.jpg', '.jpeg')) and Image is not None:
        if pytesseract is not None:
            try:
                img = Image.open(io.BytesIO(data))
                text = pytesseract.image_to_string(img)
                return text or f"(Bild {name} ohne erkennbaren Text.)"
            except Exception:
                return f"(OCR für Bild {name} fehlgeschlagen.)"
        else:
            return f"(Bild {name} angehängt; OCR nicht verfügbar in der Umgebung.)"

    # Fallback: include filename and size
    return f"(Datei {name} angehängt — Typ unbekannt oder nicht unterstützte Extraktion. Größe: {len(data)} bytes)"

# Keep chat history
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display previous messages
for m in st.session_state.messages:
    with st.chat_message(m["role"]):
        st.write(m["content"])
        render_sources(m.get("sources"))

with st.container(border=True):
    model_col, web_col, save_col = st.columns(3)
    with model_col:
        selected_model_label = st.selectbox(
            "Model-Selector",
            options=list(MODEL_CHOICES.keys()),
            index=0,
            key="model_selector",
        )
    with web_col:
        selected_web_search_label = st.selectbox(
            "Websuche",
            options=list(WEB_SEARCH_CHOICES.keys()),
            index=0,
            key="web_search_selector",
        )
    with save_col:
        selected_export_label = st.selectbox(
            "Chat speichern",
            options=list(EXPORT_CHOICES.keys()),
            index=0,
            disabled=not st.session_state.messages,
            key="export_selector",
        )

        selected_export = EXPORT_CHOICES[selected_export_label]
        if selected_export == "txt":
            st.download_button(
                "Download",
                data=build_conversation_text(st.session_state.messages).encode("utf-8"),
                file_name="check-ki-chat.txt",
                mime="text/plain",
            )
        elif selected_export == "docx":
            docx_export = build_docx_export(st.session_state.messages)
            if docx_export is not None:
                st.download_button(
                    "Download",
                    data=docx_export,
                    file_name="check-ki-chat.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                st.button("Download", disabled=True)
        elif selected_export == "pdf":
            pdf_export = build_pdf_export(st.session_state.messages)
            if pdf_export is not None:
                st.download_button(
                    "Download",
                    data=pdf_export,
                    file_name="check-ki-chat.pdf",
                    mime="application/pdf",
                )
            else:
                st.button("Download", disabled=True)

    with st.form("chat_form", clear_on_submit=True):
        prompt = st.text_input("Frage stellen:")
        submitted = st.form_submit_button("Senden")

if submitted and prompt.strip():
    prompt = prompt.strip()
    previous_messages = st.session_state.messages.copy()
    st.session_state.messages.append({"role": "user", "content": prompt})

    # -----------------------------
    # STEP 1: Vector Search
    # -----------------------------
    try:
        search_res = client.vector_stores.search(
            vector_store_id=VECTOR_STORE_ID,
            query=prompt
        )
    except Exception as e:
        st.error(f"Search error: {e}")
        st.stop()

    # -----------------------------
    # STEP 2: Extract context
    # -----------------------------
    context_text = ""

    for hit in search_res.data:
        if hasattr(hit, "content") and hit.content:
            for part in hit.content:
                # OpenAI spec: part.type == "text"
                if part.type == "text" and hasattr(part, "text"):
                    context_text += part.text + "\n\n"

    if not context_text:
        context_text = "(Kein relevanter Kontext gefunden.)"

    # -----------------------------
    # Attachments: extract text from uploaded files
    # -----------------------------
    attachments_text = ""
    if uploaded_files:
        for f in uploaded_files:
            try:
                extracted = extract_text_from_file(f)
            except Exception as e:
                extracted = f"(Fehler beim Verarbeiten von {f.name}: {e})"
            attachments_text += f"--- Datei: {f.name} ---\n{extracted}\n\n"
        if not attachments_text:
            attachments_text = "(Hochgeladene Dateien enthalten keinen extrahierbaren Text.)"

    # -----------------------------
    # STEP 3: Ask GPT
    # -----------------------------
    final_prompt = f"""
You are a helpful and knowledgeable assistant.
Use the provided context (if meaningful) to answer the user's question.
If the context does not help, answer the question normally.

BISHERIGER CHAT:
{build_conversation_text(previous_messages)}

KONTEXT:
{context_text}

DATEIEN (aus Hochladen):
{attachments_text}

FRAGE:
{prompt}

ANTWORT:
"""

    try:
        selected_model = choose_model(selected_model_label)
        selected_web_search = WEB_SEARCH_CHOICES[selected_web_search_label]
        response_options = build_response_options(selected_web_search)
        instructions = "You are a helpful assistant."
        if selected_web_search != "off":
            instructions += " Search the web when useful, cite sources inline, and prefer reliable sources."

        st.info(f"Verwendetes Modell: {selected_model}")
        response = client.responses.create(
            model=selected_model,
            instructions=instructions,
            input=final_prompt,
            **response_options
        )
        answer = response.output_text
        sources = collect_response_sources(response)
    except Exception as e:
        st.error(f"GPT error: {e}")
        st.stop()

    # -----------------------------
    # STEP 4: Show assistant answer
    # -----------------------------
    st.session_state.messages.append({
        "role": "assistant",
        "content": answer,
        "model": selected_model,
        "sources": sources,
    })
    st.rerun()
